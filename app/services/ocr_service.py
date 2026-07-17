"""
Local, offline document text-extraction service.

- Images (png/jpg/tiff/...)  -> pytesseract (local Tesseract binary)
- PDFs (text-based)          -> Apache Tika (runs a local jar, no network)
- PDFs (scanned/image-only)  -> pdf2image (local poppler) + pytesseract
- DOCX / XLSX                -> python-docx / openpyxl
- Everything else            -> Apache Tika fallback (local server)

Nothing here calls out to the internet. Tika's python wrapper will start a
local Tika server jar (downloaded once) and talk to it over 127.0.0.1.
"""
import os
from pathlib import Path

import pytesseract
from PIL import Image
from pdf2image import convert_from_path
from docx import Document as DocxDocument
import openpyxl

from app.config import settings

pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD

IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".gif"}
PDF_EXTS = {".pdf"}
DOCX_EXTS = {".docx"}
XLSX_EXTS = {".xlsx"}


def _ocr_image(path: str) -> str:
    img = Image.open(path)
    return pytesseract.image_to_string(img)


def _extract_pdf(path: str) -> str:
    """Try Tika first (fast, works for text PDFs); fall back to OCR per page."""
    text = ""
    try:
        from tika import parser as tika_parser
        parsed = tika_parser.from_file(path)
        text = (parsed.get("content") or "").strip()
    except Exception:
        text = ""

    if len(text) > 20:
        return text

    # Fallback: scanned PDF -> rasterize pages locally and OCR each one
    pages = convert_from_path(path, dpi=300)
    ocr_text = []
    for i, page in enumerate(pages):
        tmp_path = f"{path}.page{i}.png"
        page.save(tmp_path, "PNG")
        try:
            ocr_text.append(pytesseract.image_to_string(Image.open(tmp_path)))
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
    return "\n".join(ocr_text)


def _extract_docx(path: str) -> str:
    doc = DocxDocument(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_xlsx(path: str) -> str:
    wb = openpyxl.load_workbook(path, data_only=True)
    lines = []
    for sheet in wb.worksheets:
        lines.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            lines.append(" | ".join("" if v is None else str(v) for v in row))
    return "\n".join(lines)


def _extract_tika_generic(path: str) -> str:
    from tika import parser as tika_parser
    parsed = tika_parser.from_file(path)
    return (parsed.get("content") or "").strip()


def extract_text(path: str) -> str:
    """Route a document to the right local extractor based on extension."""
    ext = Path(path).suffix.lower()

    if ext in IMAGE_EXTS:
        return _ocr_image(path)
    if ext in PDF_EXTS:
        return _extract_pdf(path)
    if ext in DOCX_EXTS:
        return _extract_docx(path)
    if ext in XLSX_EXTS:
        return _extract_xlsx(path)
    # .txt, .html, .rtf, .odt, etc.
    if ext == ".txt":
        return Path(path).read_text(errors="ignore")

    return _extract_tika_generic(path)
